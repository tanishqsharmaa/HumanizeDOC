"""Tests for the DOCX parser (``pipeline.parser``)."""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from humanizedoc.backend.models import BlockType, DocumentBlock
from humanizedoc.backend.pipeline.parser import parse_docx, count_words


# ═══════════════════════════════════════════════════════════════════
# Helpers — create in-memory DOCX fixtures and save to temp files
# ═══════════════════════════════════════════════════════════════════

def _save_temp_docx(doc: Document) -> str:
    """Save a python-docx Document to a temp file and return its path."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _make_simple_doc() -> str:
    """Create a simple doc with a heading, two body paragraphs, and an empty paragraph."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    p1 = doc.add_paragraph("This is the first paragraph with enough words to be humanized.")
    run = p1.runs[0]
    run.bold = True
    doc.add_paragraph("")  # empty paragraph
    doc.add_paragraph("This is the second paragraph also with plenty of words for testing.")
    return _save_temp_docx(doc)


def _make_doc_with_table() -> str:
    """Create a doc with a 2×2 table."""
    doc = Document()
    doc.add_paragraph("Body text before the table has many words in it for sure.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Data cell alpha with enough words to be meaningful text"
    table.cell(1, 1).text = "Data cell beta with sufficient content length for testing"
    return _save_temp_docx(doc)


# ═══════════════════════════════════════════════════════════════════
# Tests
# ═══════════════════════════════════════════════════════════════════

class TestParseDocx:
    """Tests for ``parse_docx``."""

    def test_returns_list_of_document_blocks(self):
        path = _make_simple_doc()
        blocks = parse_docx(path)
        assert isinstance(blocks, list)
        assert all(isinstance(b, DocumentBlock) for b in blocks)
        Path(path).unlink(missing_ok=True)

    def test_heading_paragraph_has_correct_style(self):
        path = _make_simple_doc()
        blocks = parse_docx(path)
        heading_block = blocks[0]
        assert heading_block.style_name == "Heading 1"
        assert heading_block.text == "Introduction"
        Path(path).unlink(missing_ok=True)

    def test_empty_paragraph_preserved(self):
        path = _make_simple_doc()
        blocks = parse_docx(path)
        # The empty paragraph is the third block (index 2)
        empty_block = [b for b in blocks if b.text == ""]
        assert len(empty_block) >= 1
        Path(path).unlink(missing_ok=True)

    def test_bold_formatting_extracted(self):
        path = _make_simple_doc()
        blocks = parse_docx(path)
        # First body paragraph (index 1) has bold run
        body_block = blocks[1]
        assert body_block.formatting.bold is True
        Path(path).unlink(missing_ok=True)

    def test_table_paragraphs_have_is_table_cell(self):
        path = _make_doc_with_table()
        blocks = parse_docx(path)
        table_blocks = [b for b in blocks if b.formatting.is_table_cell]
        assert len(table_blocks) == 4  # 2×2 table = 4 cells
        Path(path).unlink(missing_ok=True)

    def test_table_block_id_format(self):
        path = _make_doc_with_table()
        blocks = parse_docx(path)
        table_blocks = [b for b in blocks if b.block_id.startswith("table_")]
        assert len(table_blocks) >= 1
        # First table block should reference row 0
        assert "row_0" in table_blocks[0].block_id
        Path(path).unlink(missing_ok=True)

    def test_invalid_file_raises_value_error(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(b"not a real docx")
        tmp.close()
        with pytest.raises(ValueError, match="Invalid|Failed"):
            parse_docx(tmp.name)
        Path(tmp.name).unlink(missing_ok=True)

    def test_nonexistent_file_raises_value_error(self):
        with pytest.raises(ValueError):
            parse_docx("/nonexistent/path/file.docx")


class TestCountWords:
    """Tests for ``count_words``."""

    def test_count_words_basic(self):
        blocks = [
            DocumentBlock(
                block_id="b_0001",
                block_type=BlockType.HUMANIZE,
                text="Hello world foo bar baz",
                paragraph_index=0,
            ),
            DocumentBlock(
                block_id="b_0002",
                block_type=BlockType.PRESERVE,
                text="This should not count",
                paragraph_index=1,
            ),
            DocumentBlock(
                block_id="b_0003",
                block_type=BlockType.HUMANIZE,
                text="One two three",
                paragraph_index=2,
            ),
        ]
        assert count_words(blocks) == 8  # 5 + 3

    def test_count_words_empty(self):
        assert count_words([]) == 0
