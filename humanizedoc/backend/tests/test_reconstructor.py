"""Tests for the DOCX reconstructor (``pipeline.reconstructor``)."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from humanizedoc.backend.models import BlockType, DocumentBlock, FormattingMeta
from humanizedoc.backend.pipeline.reconstructor import (
    parse_alignment,
    reconstruct_docx,
)


# ═══════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════

def _save_temp(doc: Document) -> str:
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _make_blocks_and_original() -> tuple[str, list[DocumentBlock]]:
    """Create a minimal original doc and corresponding blocks."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    p1 = doc.add_paragraph("This is background context for the essay.")
    run = p1.runs[0]
    run.bold = True
    run.font.name = "Arial"
    doc.add_paragraph("Another paragraph that will be humanized by the system.")
    orig_path = _save_temp(doc)

    blocks = [
        DocumentBlock(
            block_id="block_0000",
            block_type=BlockType.PRESERVE,
            text="Introduction",
            paragraph_index=0,
            style_name="Heading 1",
            formatting=FormattingMeta(bold=True),
        ),
        DocumentBlock(
            block_id="block_0001",
            block_type=BlockType.HUMANIZE,
            text="This is background context for the essay.",
            paragraph_index=1,
            style_name="Normal",
            formatting=FormattingMeta(bold=True, font_name="Arial"),
        ),
        DocumentBlock(
            block_id="block_0002",
            block_type=BlockType.HUMANIZE,
            text="Another paragraph that will be humanized by the system.",
            paragraph_index=2,
            style_name="Normal",
            formatting=FormattingMeta(),
        ),
    ]
    return orig_path, blocks


# ═══════════════════════════════════════════════════════════════════
# Tests
# ═══════════════════════════════════════════════════════════════════

class TestParseAlignment:
    def test_known_values(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert parse_alignment("LEFT") == WD_ALIGN_PARAGRAPH.LEFT
        assert parse_alignment("CENTER") == WD_ALIGN_PARAGRAPH.CENTER
        assert parse_alignment("RIGHT") == WD_ALIGN_PARAGRAPH.RIGHT
        assert parse_alignment("JUSTIFY") == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_none_returns_none(self):
        assert parse_alignment(None) is None

    def test_empty_string_returns_none(self):
        assert parse_alignment("") is None


class TestReconstructDocx:
    def test_output_has_correct_paragraph_count(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())  # placeholder

        humanized_map = {
            "block_0001": "Rewritten background text for the essay.",
            "block_0002": "Rewritten another paragraph by the LLM system.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        assert len(result.paragraphs) == 3
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_preserve_blocks_appear_verbatim(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())

        humanized_map = {
            "block_0001": "Rewritten text here.",
            "block_0002": "Another rewritten chunk.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        assert result.paragraphs[0].text == "Introduction"
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_humanize_blocks_use_humanized_text(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())

        humanized_map = {
            "block_0001": "Completely rewritten background text.",
            "block_0002": "Totally different phrasing here.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        assert result.paragraphs[1].text == "Completely rewritten background text."
        assert result.paragraphs[2].text == "Totally different phrasing here."
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_font_name_preserved(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())

        humanized_map = {
            "block_0001": "Rewritten text with Arial font.",
            "block_0002": "Another rewritten chunk.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        run = result.paragraphs[1].runs[0]
        assert run.font.name == "Arial"
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_bold_formatting_preserved(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())

        humanized_map = {
            "block_0001": "Bold text should stay bold.",
            "block_0002": "Non-bold text.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        assert result.paragraphs[1].runs[0].bold is True
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_fallback_to_original_when_not_in_map(self):
        orig_path, blocks = _make_blocks_and_original()
        out_path = _save_temp(Document())

        # Only provide humanized text for one block
        humanized_map = {
            "block_0001": "Only this block was humanized.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        # block_0002 should fall back to original text
        assert result.paragraphs[2].text == "Another paragraph that will be humanized by the system."
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)

    def test_italic_formatting_preserved(self):
        orig_path, blocks = _make_blocks_and_original()
        # Make block_0002 italic
        blocks[2].formatting.italic = True
        out_path = _save_temp(Document())

        humanized_map = {
            "block_0001": "Text one.",
            "block_0002": "Italic text.",
        }
        reconstruct_docx(orig_path, blocks, humanized_map, out_path)

        result = Document(out_path)
        assert result.paragraphs[2].runs[0].italic is True
        Path(orig_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)
