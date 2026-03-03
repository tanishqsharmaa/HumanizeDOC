"""
HumanizeDOC — DOCX Parser
~~~~~~~~~~~~~~~~~~~~~~~~~~~
Extracts every paragraph from a .docx file (including table cells and
footnotes) and returns a flat list of ``DocumentBlock`` objects with
full formatting metadata.
"""

from __future__ import annotations

import logging
from typing import List, Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from humanizedoc.backend.models import (
    BlockType,
    DocumentBlock,
    FormattingMeta,
)

logger = logging.getLogger(__name__)

# ── Alignment enum → readable string ──────────────────────────────
_ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
}


# ═══════════════════════════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════════════════════════

def _resolve_bool(value: Optional[bool], fallback: Optional[bool] = None) -> bool:
    """Return *value* if it is an explicit bool, else *fallback*, else False."""
    if value is not None:
        return bool(value)
    if fallback is not None:
        return bool(fallback)
    return False


def _detect_list_info(paragraph) -> tuple[bool, int, Optional[str]]:
    """Detect whether *paragraph* is a list item and extract level + type.

    Returns ``(is_list_item, list_level, list_type)`` where *list_type*
    is ``"bullet"``, ``"numbered"``, or ``None``.
    """
    style_name: str = getattr(paragraph.style, "name", "") or ""
    if style_name.startswith("List"):
        is_list = True
    else:
        is_list = False

    # XML-level detection via <w:numPr>
    pPr = paragraph._p.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr")) if pPr is not None else None

    list_level = 0
    list_type: Optional[str] = None

    if numPr is not None:
        is_list = True
        ilvl_el = numPr.find(qn("w:ilvl"))
        if ilvl_el is not None:
            list_level = int(ilvl_el.get(qn("w:val"), "0"))

        # Heuristic: styles containing "Bullet" → bullet, else numbered
        numId_el = numPr.find(qn("w:numId"))
        if numId_el is not None:
            if "bullet" in style_name.lower() or "list bullet" in style_name.lower():
                list_type = "bullet"
            else:
                list_type = "numbered"

    return is_list, list_level, list_type


def _extract_formatting(paragraph, *, is_table_cell: bool = False,
                         is_footnote: bool = False) -> FormattingMeta:
    """Build a ``FormattingMeta`` from *paragraph* and its runs."""
    # Alignment
    alignment_str: Optional[str] = None
    if paragraph.alignment is not None:
        alignment_str = _ALIGNMENT_MAP.get(paragraph.alignment)

    # Spacing
    pf = paragraph.paragraph_format
    space_before: Optional[int] = None
    space_after: Optional[int] = None
    line_spacing: Optional[float] = None

    if pf.space_before is not None:
        space_before = int(pf.space_before.pt)
    if pf.space_after is not None:
        space_after = int(pf.space_after.pt)
    if pf.line_spacing is not None:
        try:
            line_spacing = float(pf.line_spacing)
        except TypeError:
            line_spacing = float(pf.line_spacing.pt) if pf.line_spacing else None

    # Style-level font defaults
    style_font = getattr(paragraph.style, "font", None)
    style_font_name = getattr(style_font, "name", None) if style_font else None
    style_bold = getattr(style_font, "bold", None) if style_font else None
    style_italic = getattr(style_font, "italic", None) if style_font else None
    style_underline = getattr(style_font, "underline", None) if style_font else None

    # Run-level overrides — use first run with non-None properties
    font_name: Optional[str] = style_font_name
    font_size: Optional[int] = None
    bold = _resolve_bool(None, style_bold)
    italic = _resolve_bool(None, style_italic)
    underline = _resolve_bool(None, style_underline)

    for run in paragraph.runs:
        rf = run.font
        if rf.name is not None:
            font_name = rf.name
        if rf.size is not None:
            # python-docx returns EMU; Pt() constructor accepts Pt.
            # rf.size is in EMU — convert to half-points (1 pt = 2 hp).
            font_size = int(rf.size / Pt(0.5))
            # Alternatively: half_points = int(rf.size.pt * 2)
        if rf.bold is not None:
            bold = bool(rf.bold)
        if rf.italic is not None:
            italic = bool(rf.italic)
        if rf.underline is not None:
            underline = bool(rf.underline)
        # Once we have at least one run with concrete values, stop
        if rf.name is not None or rf.size is not None or rf.bold is not None:
            break

    # List detection
    is_list_item, list_level, list_type = _detect_list_info(paragraph)

    return FormattingMeta(
        font_name=font_name,
        font_size=font_size,
        bold=bold,
        italic=italic,
        underline=underline,
        alignment=alignment_str,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
        is_list_item=is_list_item,
        list_level=list_level,
        list_type=list_type,
        is_table_cell=is_table_cell,
        is_footnote=is_footnote,
    )


# ═══════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════

def parse_docx(file_path: str) -> List[DocumentBlock]:
    """Parse a .docx file into a flat list of ``DocumentBlock`` objects.

    Raises ``ValueError`` for invalid / corrupted files.
    """
    try:
        doc = Document(file_path)
    except (PackageNotFoundError, Exception) as exc:
        if "Package" in type(exc).__name__ or "Package" in str(exc):
            raise ValueError("Invalid or corrupted .docx file") from exc
        raise ValueError(f"Failed to parse document: {exc}") from exc

    blocks: List[DocumentBlock] = []
    global_index = 0  # unique paragraph_index across entire document

    # ── 1. Body paragraphs ────────────────────────────────────────
    try:
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style_name = getattr(para.style, "name", "Normal") or "Normal"
            formatting = _extract_formatting(para)

            blocks.append(DocumentBlock(
                block_id=f"block_{global_index:04d}",
                block_type=BlockType.HUMANIZE,  # classifier sets final type
                text=text,
                paragraph_index=global_index,
                style_name=style_name,
                formatting=formatting,
            ))
            global_index += 1

        # ── 2. Table paragraphs ───────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        style_name = getattr(para.style, "name", "Normal") or "Normal"
                        formatting = _extract_formatting(
                            para, is_table_cell=True,
                        )

                        blocks.append(DocumentBlock(
                            block_id=f"table_{t_idx}_row_{r_idx}_col_{c_idx}_para_{p_idx}",
                            block_type=BlockType.HUMANIZE,
                            text=text,
                            paragraph_index=global_index,
                            style_name=style_name,
                            formatting=formatting,
                        ))
                        global_index += 1

        # ── 3. Footnote paragraphs ────────────────────────────────
        try:
            footnotes_part = doc.part.footnotes_part
            if footnotes_part is not None:
                footnotes_el = footnotes_part.element
                for fn_el in footnotes_el.findall(qn("w:footnote")):
                    fn_id = fn_el.get(qn("w:id"), "0")
                    # Skip the built-in separator footnotes (ids 0 and -1)
                    if fn_id in ("0", "-1"):
                        continue
                    for p_idx, p_el in enumerate(fn_el.findall(qn("w:p"))):
                        # Reconstruct text from <w:t> elements
                        text_parts = [
                            t.text or ""
                            for t in p_el.iter(qn("w:t"))
                        ]
                        text = "".join(text_parts).strip()

                        blocks.append(DocumentBlock(
                            block_id=f"footnote_{fn_id}_para_{p_idx}",
                            block_type=BlockType.HUMANIZE,
                            text=text,
                            paragraph_index=global_index,
                            style_name="Footnote Text",
                            formatting=FormattingMeta(is_footnote=True),
                        ))
                        global_index += 1
        except AttributeError:
            # Document has no footnotes_part — that's fine
            logger.debug("No footnotes found in document.")

    except Exception as exc:
        raise ValueError(f"Failed to parse document: {exc}") from exc

    logger.info("Parsed %d blocks from %s", len(blocks), file_path)
    return blocks


def count_words(blocks: List[DocumentBlock]) -> int:
    """Return the total word count across all HUMANIZE-type blocks."""
    total = 0
    for block in blocks:
        if block.block_type == BlockType.HUMANIZE:
            total += len(block.text.split())
    return total
