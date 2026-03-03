"""
HumanizeDOC — DOCX Reconstructor
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Rebuilds a new .docx file using humanized text while preserving the
original document's formatting, styles, page layout, and table structure.
"""

from __future__ import annotations

import copy
import logging
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Emu

from humanizedoc.backend.models import BlockType, DocumentBlock

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# Alignment helpers
# ═══════════════════════════════════════════════════════════════════

_ALIGNMENT_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def parse_alignment(alignment_str: Optional[str]) -> Optional[int]:
    """Convert a string alignment back to a python-docx enum value."""
    if not alignment_str:
        return None
    return _ALIGNMENT_MAP.get(alignment_str.upper())


# ═══════════════════════════════════════════════════════════════════
# Section property copier
# ═══════════════════════════════════════════════════════════════════

def copy_section_properties(
    original_doc: Document,
    out_doc: Document,
) -> None:
    """Copy page size, margins, and orientation from *original_doc* to *out_doc*."""
    if not original_doc.sections or not out_doc.sections:
        return

    for orig_sec, out_sec in zip(original_doc.sections, out_doc.sections):
        out_sec.page_width = orig_sec.page_width
        out_sec.page_height = orig_sec.page_height
        out_sec.left_margin = orig_sec.left_margin
        out_sec.right_margin = orig_sec.right_margin
        out_sec.top_margin = orig_sec.top_margin
        out_sec.bottom_margin = orig_sec.bottom_margin
        out_sec.gutter = orig_sec.gutter
        out_sec.orientation = orig_sec.orientation

        # Copy header / footer XML if present
        _copy_header_footer(orig_sec, out_sec)


def _copy_header_footer(orig_section, out_section) -> None:
    """Deep-copy header and footer XML elements between sections."""
    try:
        if orig_section.header and orig_section.header.is_linked_to_previous is False:
            for para in orig_section.header.paragraphs:
                new_para = out_section.header.add_paragraph()
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
        if orig_section.footer and orig_section.footer.is_linked_to_previous is False:
            for para in orig_section.footer.paragraphs:
                new_para = out_section.footer.add_paragraph()
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
    except Exception:
        logger.debug("Could not copy headers/footers — continuing.")


# ═══════════════════════════════════════════════════════════════════
# Paragraph formatting application
# ═══════════════════════════════════════════════════════════════════

def _apply_paragraph_formatting(para, block: DocumentBlock) -> None:
    """Apply paragraph-level formatting from *block.formatting* to *para*."""
    fmt = block.formatting
    pf = para.paragraph_format

    if fmt.space_before is not None:
        pf.space_before = Pt(fmt.space_before)
    if fmt.space_after is not None:
        pf.space_after = Pt(fmt.space_after)
    if fmt.line_spacing is not None:
        pf.line_spacing = fmt.line_spacing

    alignment = parse_alignment(fmt.alignment)
    if alignment is not None:
        para.alignment = alignment


def _apply_run_formatting(run, block: DocumentBlock) -> None:
    """Apply character-level formatting from *block.formatting* to *run*."""
    fmt = block.formatting

    run.bold = fmt.bold
    run.italic = fmt.italic
    run.underline = fmt.underline

    if fmt.font_name:
        run.font.name = fmt.font_name
    if fmt.font_size:
        # font_size is stored in half-points → divide by 2 for Pt
        run.font.size = Pt(fmt.font_size / 2)


# ═══════════════════════════════════════════════════════════════════
# Style helpers
# ═══════════════════════════════════════════════════════════════════

def _copy_styles_xml(original_doc: Document, out_doc: Document) -> None:
    """Copy the entire styles XML from the original document."""
    try:
        orig_styles_element = original_doc.styles.element
        out_styles_part = out_doc.part.styles_part
        if out_styles_part is not None and orig_styles_element is not None:
            out_styles_part._element = copy.deepcopy(orig_styles_element)
    except Exception:
        logger.debug("Could not copy styles XML — using defaults.")


def _safe_apply_style(para, style_name: str, out_doc: Document) -> None:
    """Set *para.style* to *style_name*, falling back to Normal."""
    try:
        para.style = out_doc.styles[style_name]
    except KeyError:
        try:
            para.style = out_doc.styles["Normal"]
        except KeyError:
            pass


# ═══════════════════════════════════════════════════════════════════
# Main reconstruction
# ═══════════════════════════════════════════════════════════════════

def reconstruct_docx(
    original_path: str,
    blocks: List[DocumentBlock],
    humanized_block_map: Dict[str, str],
    output_path: str,
) -> None:
    """Rebuild a DOCX from *blocks* with humanized text + original formatting.

    Parameters
    ----------
    original_path:
        Path to the original uploaded .docx file.
    blocks:
        Full, ordered list of ``DocumentBlock`` objects (HUMANIZE + PRESERVE).
    humanized_block_map:
        Mapping ``block_id → humanized_text`` for HUMANIZE blocks.
    output_path:
        Where to save the reconstructed .docx.
    """
    original_doc = Document(original_path)
    out_doc = Document()

    # ── Copy document-level properties ────────────────────────────
    _copy_styles_xml(original_doc, out_doc)
    copy_section_properties(original_doc, out_doc)

    # ── Remove the default empty paragraph python-docx adds ──────
    if out_doc.paragraphs:
        _remove_paragraph(out_doc.paragraphs[0])

    # ── Separate body blocks from table / footnote blocks ─────────
    body_blocks = [
        b for b in blocks
        if not b.block_id.startswith("table_") and not b.block_id.startswith("footnote_")
    ]
    table_blocks = [b for b in blocks if b.block_id.startswith("table_")]

    # ── Rebuild body paragraphs ───────────────────────────────────
    for block in _sorted_blocks(body_blocks):
        final_text = _get_final_text(block, humanized_block_map)

        new_para = out_doc.add_paragraph()
        _safe_apply_style(new_para, block.style_name, out_doc)
        _apply_paragraph_formatting(new_para, block)

        run = new_para.add_run(final_text)
        _apply_run_formatting(run, block)

    # ── Rebuild tables ────────────────────────────────────────────
    _reconstruct_tables(original_doc, out_doc, table_blocks, humanized_block_map)

    # ── Save ──────────────────────────────────────────────────────
    out_doc.save(output_path)
    logger.info("Reconstructed document saved to %s", output_path)


# ═══════════════════════════════════════════════════════════════════
# Table reconstruction
# ═══════════════════════════════════════════════════════════════════

def _reconstruct_tables(
    original_doc: Document,
    out_doc: Document,
    table_blocks: List[DocumentBlock],
    humanized_block_map: Dict[str, str],
) -> None:
    """Rebuild tables from the original doc, applying humanized text."""
    if not original_doc.tables:
        return

    # Index table blocks by (table_idx, row_idx, col_idx, para_idx)
    tb_map: Dict[tuple, DocumentBlock] = {}
    for b in table_blocks:
        key = _parse_table_block_id(b.block_id)
        if key is not None:
            tb_map[key] = b

    for t_idx, orig_table in enumerate(original_doc.tables):
        rows = len(orig_table.rows)
        cols = len(orig_table.columns)
        new_table = out_doc.add_table(rows=rows, cols=cols)

        # Copy table style if available
        try:
            new_table.style = orig_table.style
        except Exception:
            pass

        for r_idx, row in enumerate(orig_table.rows):
            for c_idx, cell in enumerate(row.cells):
                new_cell = new_table.cell(r_idx, c_idx)
                # Clear default paragraph
                if new_cell.paragraphs:
                    for i, p in enumerate(new_cell.paragraphs):
                        if i == 0:
                            continue  # keep first, reuse it
                        _remove_paragraph(p)

                for p_idx, orig_para in enumerate(cell.paragraphs):
                    key = (t_idx, r_idx, c_idx, p_idx)
                    block = tb_map.get(key)

                    if p_idx == 0:
                        new_para = new_cell.paragraphs[0]
                    else:
                        new_para = new_cell.add_paragraph()

                    if block is not None:
                        final_text = _get_final_text(block, humanized_block_map)
                        _safe_apply_style(new_para, block.style_name, out_doc)
                        _apply_paragraph_formatting(new_para, block)
                        run = new_para.add_run(final_text)
                        _apply_run_formatting(run, block)
                    else:
                        # No block found — copy original text verbatim
                        new_para.add_run(orig_para.text)


def _parse_table_block_id(block_id: str) -> Optional[tuple]:
    """Parse ``table_{t}_row_{r}_col_{c}_para_{p}`` → ``(t, r, c, p)``."""
    parts = block_id.split("_")
    try:
        t = int(parts[parts.index("table") + 1])
        r = int(parts[parts.index("row") + 1])
        c = int(parts[parts.index("col") + 1])
        p = int(parts[parts.index("para") + 1])
        return (t, r, c, p)
    except (ValueError, IndexError):
        return None


# ═══════════════════════════════════════════════════════════════════
# Utility
# ═══════════════════════════════════════════════════════════════════

def _get_final_text(
    block: DocumentBlock,
    humanized_block_map: Dict[str, str],
) -> str:
    """Return the text to write: humanized if available, else original."""
    if block.block_type == BlockType.PRESERVE:
        return block.text
    return humanized_block_map.get(block.block_id, block.text)


def _sorted_blocks(blocks: List[DocumentBlock]) -> List[DocumentBlock]:
    """Return blocks sorted by paragraph_index."""
    return sorted(blocks, key=lambda b: b.paragraph_index)


def _remove_paragraph(para) -> None:
    """Remove a paragraph element from its parent XML tree."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
